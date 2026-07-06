from docx import Document
from parser3.document import DocumentReader
from parser3.headings import SectionClassifier


def test_document_reader_adds_generated_heading_numbers(tmp_path):
    path = tmp_path / "headings.docx"
    doc = Document()
    doc.add_paragraph("Kapitel", style="Heading 1")
    doc.add_paragraph("Underkapitel", style="Heading 2")
    doc.save(path)

    blocks = DocumentReader().read(path)

    assert blocks[0].text == "1 § Kapitel"
    assert blocks[1].text == "1.1 § Underkapitel"
    assert SectionClassifier().classify(blocks[0].text).number == "1"
    assert SectionClassifier().classify(blocks[1].text).number == "1.1"
