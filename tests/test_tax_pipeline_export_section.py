from docx import Document
from parser3.pipeline import TaxPipeline


def test_tax_pipeline_exports_tax_row_outside_chapter_1(tmp_path):
    path = tmp_path / "taxa_export.docx"
    doc = Document()
    doc.add_paragraph("Avgifter och tjänster", style="Heading 1")
    doc.add_paragraph("Grundavgift", style="Heading 2")

    # HeadingNumberer makes these 1 and 1.1, so use an explicit section heading
    # to simulate a real non-chapter-1 tax paragraph.
    doc.add_paragraph("2.1 § Grundavgift")
    doc.add_paragraph("Fritidshus XX kr")
    doc.save(path)

    result = TaxPipeline().run(path)

    assert result.tax_rows
    assert result.tax_rows[0].name == "Fritidshus"
