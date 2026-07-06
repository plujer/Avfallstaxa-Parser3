from docx import Document
from parser3.pipeline import TaxPipeline


def test_tax_pipeline_runs_single_official_flow(tmp_path):
    path = tmp_path / "taxa.docx"
    doc = Document()
    doc.add_paragraph("Avgifter", style="Heading 1")
    doc.add_paragraph("Grundavgift", style="Heading 2")
    doc.add_paragraph("Fritidshus XX kr")
    doc.save(path)

    result = TaxPipeline().run(path)

    assert result.blocks
    assert result.semantic_rows

    # Chapter 1 is definitions/legal text in the real parser rules and should
    # not export tax rows. This test verifies the official pipeline runs.
    assert result.tax_rows == []
