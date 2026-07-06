"""Word document reader for Parser 3.0.

This module reads paragraphs and native Word tables into a common sequence of
DocumentBlock objects. It also reconstructs automatic Word heading numbering,
because python-docx does not include generated heading numbers in paragraph.text.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document

from parser3.document.heading_numberer import HeadingNumberer
from parser3.utils.exceptions import DocumentReadError


@dataclass
class DocumentBlock:
    order: int
    kind: str
    text: str
    style: str = ""
    rows: list[list[str]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentReader:
    """Read a .docx file into ordered document blocks."""

    def read(self, path: str | Path) -> list[DocumentBlock]:
        doc_path = Path(path)
        if not doc_path.exists():
            raise DocumentReadError(f"Document not found: {doc_path}")
        if doc_path.suffix.lower() != ".docx":
            raise DocumentReadError(f"Only .docx is supported in this block: {doc_path}")

        try:
            document = Document(str(doc_path))
        except Exception as exc:
            raise DocumentReadError(f"Could not open Word document: {doc_path}") from exc

        blocks: list[DocumentBlock] = []
        order = 0
        heading_numberer = HeadingNumberer()

        for child in document.element.body:
            tag = child.tag.lower()

            if tag.endswith("}p"):
                paragraph = self._paragraph_from_element(document, child)
                raw_text = self._clean(paragraph.text)
                style = paragraph.style.name if paragraph.style else ""

                if raw_text:
                    text, metadata = heading_numberer.prefix_heading(raw_text, style)
                    blocks.append(
                        DocumentBlock(
                            order=order,
                            kind="paragraph",
                            text=text,
                            style=style,
                            metadata=metadata,
                        )
                    )
                    order += 1

            elif tag.endswith("}tbl"):
                table = self._table_from_element(document, child)
                rows = self._extract_table_rows(table)
                flat_text = " | ".join(" | ".join(row) for row in rows)
                blocks.append(DocumentBlock(order=order, kind="table", text=flat_text, rows=rows))
                order += 1

        return blocks

    def _paragraph_from_element(self, document: Document, element: Any):
        for p in document.paragraphs:
            if p._element is element:
                return p
        raise DocumentReadError("Paragraph element could not be resolved.")

    def _table_from_element(self, document: Document, element: Any):
        for t in document.tables:
            if t._element is element:
                return t
        raise DocumentReadError("Table element could not be resolved.")

    def _extract_table_rows(self, table: Any) -> list[list[str]]:
        rows: list[list[str]] = []
        for row in table.rows:
            values = [self._clean(cell.text) for cell in row.cells]
            if any(values):
                rows.append(values)
        return rows

    @staticmethod
    def _clean(text: str) -> str:
        return " ".join((text or "").replace("\xa0", " ").split())
